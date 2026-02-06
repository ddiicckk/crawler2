#!/usr/bin/env python3
import argparse
import base64
import re
import shutil
import time
import tempfile
from io import BytesIO
from pathlib import Path
from datetime import datetime
from urllib.parse import unquote, urljoin, urlparse

from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError
from playwright._impl._errors import Error as PlaywrightError

from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from docx.shared import Inches, Pt


DEFAULT_URL = "https://myservice.lloyds.com/now/nav/ui/classic/params/target/kb_view.do%3Fsysparm_article%3DKB0010611"


# -------------------------
# Utilities
# -------------------------

def stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def safe_filename(name: str, default: str = "page") -> str:
    name = (name or "").strip() or default
    name = re.sub(r"[^\w\-\.]+", "_", name)
    return name[:180]


def looks_like_login(url: str) -> bool:
    u = (url or "").lower()
    return any(s in u for s in ("login", "sso", "saml", "auth", "okta", "adfs", "signin", "microsoftonline.com"))


def extract_kb_number(text: str) -> str | None:
    m = re.search(r"(KB\d+)", text or "")
    return m.group(1) if m else None


def decode_target_to_direct_url(url: str) -> str:
    """
    Converts:
      .../target/kb_view.do%3Fsysparm_article%3DKB0010611
    into:
      https://<host>/kb_view.do?sysparm_article=KB0010611
    """
    m = re.search(r"/target/([^?]+)", url)
    if not m:
        return url
    encoded = m.group(1)
    decoded = unquote(encoded)
    if decoded.startswith("http://") or decoded.startswith("https://"):
        return decoded
    host = re.match(r"^(https?://[^/]+)", url)
    if not host:
        return url
    return f"{host.group(1)}/{decoded.lstrip('/')}"


def attach_debug(page):
    page.on("console", lambda msg: print(f"[console] {msg.type}: {msg.text}"))
    page.on("pageerror", lambda err: print(f"[pageerror] {err}"))
    page.on("requestfailed", lambda req: print(f"[requestfailed] {req.url} -> {req.failure}"))


def get_body_text(page) -> str:
    try:
        return page.evaluate("() => (document.body && document.body.innerText || '').trim()")
    except Exception:
        return ""


def wait_for_text_length(page, min_chars: int, timeout_ms: int) -> int:
    deadline = time.time() + timeout_ms / 1000
    last_len = 0
    while time.time() < deadline:
        txt = get_body_text(page)
        last_len = len(txt)
        if last_len >= min_chars:
            return last_len
        time.sleep(0.25)
    return last_len


def goto_with_sso_retry(page, url: str, timeout_ms: int, headed: bool, max_attempts: int = 3) -> None:
    for attempt in range(1, max_attempts + 1):
        try:
            page.goto(url, wait_until="domcontentloaded", timeout=timeout_ms)
            return
        except PlaywrightTimeoutError:
            print(f"⚠️ goto timeout (attempt {attempt}/{max_attempts}). Continuing...")
            return
        except PlaywrightError as e:
            msg = str(e)
            if "is interrupted by another navigation" in msg:
                print(f"⚠️ Navigation interrupted by SSO (attempt {attempt}/{max_attempts}). URL now: {page.url}")
                try:
                    page.wait_for_load_state("domcontentloaded", timeout=30_000)
                except Exception:
                    pass

                if looks_like_login(page.url):
                    if not headed:
                        raise RuntimeError("SSO requires interaction but you are running headless. Re-run with --headed.")
                    input("✅ Complete SSO in the browser, then press ENTER to retry...")
                continue
            raise


# -------------------------
# Main content extraction
# -------------------------

USELESS_TAGS = {"script", "style", "noscript", "svg", "canvas"}
USELESS_SELECTORS = [
    "header", "footer", "nav",
    "[role='navigation']", "[role='banner']", "[role='contentinfo']",
    ".navbar", ".nav", ".navigation", ".header", ".footer",
    ".sidebar", ".side-nav", ".toc", ".breadcrumbs",
]


def soup_remove_useless(soup: BeautifulSoup) -> None:
    for t in list(soup.find_all(list(USELESS_TAGS))):
        t.decompose()
    for sel in USELESS_SELECTORS:
        for t in soup.select(sel):
            t.decompose()


def pick_best_container(soup: BeautifulSoup) -> Tag:
    candidate_selectors = [
        "article", "main",
        "#kb_article", ".kb-article", ".kb-article-content", ".kb-view",
        ".sn-kb-article", ".knowledge-article",
        "[id*='kb']", "[class*='kb']",
        "[class*='article']", "[class*='content']",
    ]

    best_node = None
    best_len = 0

    for sel in candidate_selectors:
        for node in soup.select(sel):
            txt = node.get_text("\n", strip=True)
            ln = len(txt)
            if ln > best_len:
                best_len = ln
                best_node = node

    if best_node and best_len >= 400:
        return best_node

    # fallback: largest meaningful block
    for node in soup.find_all(["div", "section", "article", "main"], recursive=True):
        txt = node.get_text("\n", strip=True)
        ln = len(txt)
        if ln > best_len:
            best_len = ln
            best_node = node

    return best_node if best_node else (soup.body or soup)


def extract_main_container(html: str) -> tuple[str, Tag]:
    """
    Returns (title, container_tag) where container_tag contains main content (including images).
    Uses BS4 heuristics; optional trafilatura just for sanity but we still keep the DOM for images.
    """
    soup = BeautifulSoup(html, "lxml")
    title = soup.title.get_text(strip=True) if soup.title else ""
    soup_remove_useless(soup)
    container = pick_best_container(soup)
    return title, container


# -------------------------
# Image downloading (auth-aware)
# -------------------------

def is_data_image(src: str) -> bool:
    return src.startswith("data:image/")


def decode_data_image(src: str) -> tuple[bytes, str]:
    # data:image/png;base64,....
    header, b64 = src.split(",", 1)
    # try to infer ext
    m = re.search(r"data:image/([^;]+);base64", header, re.IGNORECASE)
    ext = (m.group(1).lower() if m else "png").replace("jpeg", "jpg")
    return base64.b64decode(b64), ext


def guess_ext_from_content_type(ct: str | None) -> str:
    if not ct:
        return "png"
    ct = ct.lower()
    if "jpeg" in ct:
        return "jpg"
    if "png" in ct:
        return "png"
    if "gif" in ct:
        return "gif"
    if "webp" in ct:
        return "webp"
    if "bmp" in ct:
        return "bmp"
    if "tiff" in ct:
        return "tif"
    return "png"


def fetch_image_bytes(page, context, src: str, base_url: str) -> tuple[bytes | None, str | None]:
    """
    Returns (bytes, ext) or (None, None) if cannot fetch.
    Handles:
      - data:image/... base64
      - normal http(s) urls (auth via context.request)
      - blob: urls (fallback using fetch in-page)
    """
    if not src:
        return None, None

    src = src.strip()

    # 1) data URI
    if is_data_image(src):
        bts, ext = decode_data_image(src)
        return bts, ext

    # 2) blob URL: must be fetched in the page context
    if src.startswith("blob:"):
        try:
            # fetch blob in page and return base64
            b64 = page.evaluate(
                """async (u) => {
                    const r = await fetch(u);
                    const b = await r.blob();
                    const a = new Uint8Array(await b.arrayBuffer());
                    let s = '';
                    for (let i=0; i<a.length; i++) s += String.fromCharCode(a[i]);
                    return btoa(s);
                }""",
                src
            )
            bts = base64.b64decode(b64)
            # blob doesn't carry extension well; default to png
            return bts, "png"
        except Exception:
            return None, None

    # 3) normal URL: resolve relative
    full_url = urljoin(base_url, src)

    # Ignore non-http(s)
    scheme = urlparse(full_url).scheme.lower()
    if scheme not in ("http", "https"):
        return None, None

    try:
        resp = context.request.get(full_url, timeout=60_000)
        if not resp.ok:
            return None, None
        ct = resp.headers.get("content-type")
        ext = guess_ext_from_content_type(ct)
        return resp.body(), ext
    except Exception:
        return None, None


# -------------------------
# HTML -> DOCX with images
# -------------------------

def add_paragraph_text(doc: Document, text: str):
    text = re.sub(r"\s+", " ", (text or "")).strip()
    if text:
        doc.add_paragraph(text)


def add_heading(doc: Document, text: str, level: int):
    text = re.sub(r"\s+", " ", (text or "")).strip()
    if text:
        doc.add_heading(text, level=level)


def add_list_item(doc: Document, text: str, ordered: bool):
    text = re.sub(r"\s+", " ", (text or "")).strip()
    if not text:
        return
    style = "List Number" if ordered else "List Bullet"
    doc.add_paragraph(text, style=style)


def container_to_docx(doc: Document, container: Tag, page, context, base_url: str,
                      max_image_width_in: float = 6.0):
    """
    Walk through container children in order and add to doc.
    Supports: headings, paragraphs, lists, images.
    """
    temp_dir = Path(tempfile.mkdtemp(prefix="kb_docx_imgs_"))

    def handle_node(node: Tag | NavigableString, list_mode: str | None = None):
        if isinstance(node, NavigableString):
            return

        if not isinstance(node, Tag):
            return

        name = node.name.lower()

        # skip invisible/empty wrappers
        if name in ("script", "style", "noscript"):
            return

        # Headings
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = min(4, int(name[1]))  # keep Word heading levels reasonable
            add_heading(doc, node.get_text(" ", strip=True), level=level)
            return

        # Paragraph
        if name == "p":
            # if paragraph contains images inline, add text then image(s)
            txt = node.get_text(" ", strip=True)
            if txt:
                add_paragraph_text(doc, txt)
            for img in node.find_all("img"):
                handle_node(img)
            return

        # Lists
        if name in ("ul", "ol"):
            ordered = (name == "ol")
            for li in node.find_all("li", recursive=False):
                li_text = li.get_text(" ", strip=True)
                add_list_item(doc, li_text, ordered=ordered)
                # images inside list items
                for img in li.find_all("img"):
                    handle_node(img)
            return

        # Image
        if name == "img":
            src = node.get("src") or node.get("data-src") or ""
            alt = node.get("alt") or ""
            bts, ext = fetch_image_bytes(page, context, src, base_url=base_url)

            if not bts:
                # fallback: try screenshotting the element itself
                try:
                    # element screenshot only works if we can locate it; use CSS selector if possible
                    # As a fallback, skip silently
                    if alt:
                        doc.add_paragraph(f"[Image not downloaded: {alt}]")
                    else:
                        doc.add_paragraph("[Image not downloaded]")
                except Exception:
                    pass
                return

            img_path = temp_dir / f"img_{len(list(temp_dir.iterdir())) + 1}.{ext}"
            img_path.write_bytes(bts)

            # Optional caption with alt text
            if alt:
                doc.add_paragraph(alt)

            try:
                doc.add_picture(str(img_path), width=Inches(max_image_width_in))
            except Exception:
                # if docx can't handle format (e.g., webp), skip with note
                doc.add_paragraph(f"[Image format not supported in Word: .{ext}]")
            return

        # Tables (optional): keep simple text representation
        if name == "table":
            txt = node.get_text("\n", strip=True)
            if txt:
                doc.add_paragraph(txt)
            return

        # Generic containers: recurse
        # For div/section/etc, we recurse into direct children to preserve order
        for child in node.children:
            if isinstance(child, Tag):
                handle_node(child)
            # ignore raw strings here to avoid adding lots of chrome whitespace

    # Walk direct children of the container to preserve flow
    for child in container.children:
        if isinstance(child, Tag):
            handle_node(child)

    # cleanup temp dir (best effort)
    try:
        shutil.rmtree(temp_dir, ignore_errors=True)
    except Exception:
        pass


def build_docx(title: str, source_url: str, container: Tag, page, context,
               output_path: Path):
    doc = Document()

    # Styling
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title or "Knowledge Article", level=1)
    doc.add_paragraph(f"Source: {source_url}")
    doc.add_paragraph("")

    container_to_docx(
        doc,
        container=container,
        page=page,
        context=context,
        base_url=source_url,
        max_image_width_in=6.0
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# -------------------------
# Main
# -------------------------

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--url", default=DEFAULT_URL)
    ap.add_argument("--outdir", default="downloads")
    ap.add_argument("--profile-dir", default="pw_profile")
    ap.add_argument("--relogin", action="store_true")
    ap.add_argument("--headed", action="store_true")
    ap.add_argument("--timeout", type=int, default=240_000)
    ap.add_argument("--min-chars", type=int, default=800)
    ap.add_argument("--name", default="")
    ap.add_argument("--skip-direct", action="store_true")
    args = ap.parse_args()

    out_dir = Path(args.outdir)
    out_dir.mkdir(parents=True, exist_ok=True)

    profile_dir = Path(args.profile_dir)
    if args.relogin and profile_dir.exists():
        print(f"🧹 Removing profile dir for relogin: {profile_dir}")
        shutil.rmtree(profile_dir, ignore_errors=True)

    direct_url = decode_target_to_direct_url(args.url)
    kb = extract_kb_number(direct_url) or extract_kb_number(args.url) or "KB"
    base_name = safe_filename(args.name if args.name.strip() else kb)
    base = f"{base_name}_{stamp()}"

    docx_path = out_dir / f"{base}.docx"
    html_path = out_dir / f"{base}.html"
    png_path = out_dir / f"{base}.png"

    with sync_playwright() as p:
        context = p.chromium.launch_persistent_context(
            user_data_dir=str(profile_dir),
            headless=(not args.headed),
            viewport={"width": 1400, "height": 900},
        )
        page = context.new_page()
        page.set_default_timeout(args.timeout)
        attach_debug(page)

        # Open initial URL
        print("➡️ Opening initial URL...")
        goto_with_sso_retry(page, args.url, timeout_ms=args.timeout, headed=args.headed)

        # Manual SSO if needed
        if looks_like_login(page.url):
            if not args.headed:
                context.close()
                raise RuntimeError("SSO requires interaction. Re-run with --headed.")
            input("✅ Complete SSO in the browser, then press ENTER here...")

        # Navigate to direct KB URL unless skipped
        if not args.skip_direct:
            print("➡️ Navigating to direct KB URL...")
            goto_with_sso_retry(page, direct_url, timeout_ms=args.timeout, headed=args.headed)

        # Wait for content
        print(f"⏳ Waiting for body text length >= {args.min_chars}")
        length = wait_for_text_length(page, min_chars=args.min_chars, timeout_ms=args.timeout)
        print(f"   Observed text length: {length}")

        # Save screenshot + HTML (debug / archive)
        try:
            page.screenshot(path=str(png_path), full_page=True)
            print(f"✅ Screenshot: {png_path}")
        except Exception as e:
            print(f"⚠️ Screenshot failed: {e}")

        html = page.content()
        html_path.write_text(html, encoding="utf-8")
        print(f"✅ HTML: {html_path}")

        # Extract main container + build Word doc (with images)
        title, container = extract_main_container(html)
        print("💾 Building Word document (with images)...")
        build_docx(title=title, source_url=page.url, container=container, page=page, context=context, output_path=docx_path)
        print(f"✅ Word saved: {docx_path}")

        context.close()
        print("\nDone.")


if __name__ == "__main__":
    main()
