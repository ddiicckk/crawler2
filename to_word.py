#!/usr/bin/env python3
import argparse
import re
import shutil
import time
from pathlib import Path
from datetime import datetime
from urllib.parse import unquote

from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError
from playwright._impl._errors import Error as PlaywrightError

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt


DEFAULT_URL = "https://myservice.lloyds.com/now/nav/ui/classic/params/target/kb_view.do%3Fsysparm_article%3DKB0010611"


def stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def safe_filename(name: str, default: str = "page") -> str:
    name = (name or "").strip() or default
    name = re.sub(r"[^\w\-\.]+", "_", name)
    return name[:180]


def looks_like_login(url: str) -> bool:
    u = (url or "").lower()
    return any(s in u for s in ("login", "sso", "saml", "auth", "okta", "adfs", "signin", "microsoftonline.com"))


def extract_kb_number(text: str) -> str | None:
    m = re.search(r"(KB\d+)", text or "")
    return m.group(1) if m else None


def decode_target_to_direct_url(url: str) -> str:
    m = re.search(r"/target/([^?]+)", url)
    if not m:
        return url
    encoded = m.group(1)
    decoded = unquote(encoded)
    if decoded.startswith("http://") or decoded.startswith("https://"):
        return decoded
    host = re.match(r"^(https?://[^/]+)", url)
    if not host:
        return url
    return f"{host.group(1)}/{decoded.lstrip('/')}"


def attach_debug(page):
    page.on("console", lambda msg: print(f"[console] {msg.type}: {msg.text}"))
    page.on("pageerror", lambda err: print(f"[pageerror] {err}"))
    page.on("requestfailed", lambda req: print(f"[requestfailed] {req.url} -> {req.failure}"))


def get_body_text(page) -> str:
    try:
        return page.evaluate("() => (document.body && document.body.innerText || '').trim()")
    except Exception:
        return ""


def wait_for_text_length(page, min_chars: int, timeout_ms: int) -> int:
    deadline = time.time() + timeout_ms / 1000
    last_len = 0
    while time.time() < deadline:
        txt = get_body_text(page)
        last_len = len(txt)
        if last_len >= min_chars:
            return last_len
        time.sleep(0.25)
    return last_len


def goto_with_sso_retry(page, url: str, timeout_ms: int, headed: bool, max_attempts: int = 3) -> None:
    for attempt in range(1, max_attempts + 1):
        try:
            page.goto(url, wait_until="domcontentloaded", timeout=timeout_ms)
            return
        except PlaywrightTimeoutError:
            print(f"⚠️ goto timeout (attempt {attempt}/{max_attempts}). Continuing...")
            return
        except PlaywrightError as e:
            msg = str(e)
            if "is interrupted by another navigation" in msg:
                print(f"⚠️ Navigation interrupted by SSO (attempt {attempt}/{max_attempts}). URL now: {page.url}")
                try:
                    page.wait_for_load_state("domcontentloaded", timeout=30_000)
                except Exception:
                    pass

                if looks_like_login(page.url):
                    if not headed:
                        raise RuntimeError("SSO requires interaction but you are running headless. Re-run with --headed.")
                    input("✅ Complete SSO in the browser, then press ENTER to retry...")
                continue
            raise


# -----------------------------
# MAIN CONTENT EXTRACTION
# -----------------------------

USELESS_TAGS = {"script", "style", "noscript", "svg", "canvas", "iframe"}
USELESS_ROLES = {"navigation", "banner", "contentinfo"}  # nav/header/footer semantics
USELESS_SELECTORS = [
    "header", "footer", "nav",
    "[role='navigation']", "[role='banner']", "[role='contentinfo']",
    ".navbar", ".nav", ".navigation", ".header", ".footer",
    ".sidebar", ".side-nav", ".toc", ".breadcrumbs",
    ".sn-viewport", ".sn-polaris-layout",  # sometimes large wrappers
]


def normalize_lines(text: str) -> list[str]:
    # Keep meaningful lines, collapse whitespace
    lines = []
    for raw in text.splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if not line:
            continue
        # drop obvious “chrome” lines (tweak if needed)
        if line.lower() in {"search", "home"}:
            continue
        lines.append(line)
    return lines


def soup_remove_useless(soup: BeautifulSoup) -> None:
    # Remove tag types
    for t in list(soup.find_all(list(USELESS_TAGS))):
        t.decompose()

    # Remove typical chrome areas
    for sel in USELESS_SELECTORS:
        for t in soup.select(sel):
            t.decompose()

    # Remove by role
    for role in USELESS_ROLES:
        for t in soup.select(f"[role='{role}']"):
            t.decompose()


def pick_best_container(soup: BeautifulSoup) -> BeautifulSoup:
    """
    Try common main-content containers first; fallback to the biggest text block.
    """
    # Common candidates for KB/article pages (ServiceNow & generic)
    candidate_selectors = [
        "article",
        "main",
        "#kb_article",
        ".kb-article",
        ".kb-article-content",
        ".kb-view",
        ".sn-kb-article",
        ".knowledge-article",
        "[data-testid*='article']",
        "[id*='kb']",
        "[class*='kb']",
        "[class*='article']",
        "[class*='content']",
    ]

    best_node = None
    best_len = 0

    # 1) Try known selectors
    for sel in candidate_selectors:
        for node in soup.select(sel):
            txt = node.get_text("\n", strip=True)
            ln = len(txt)
            if ln > best_len:
                best_len = ln
                best_node = node

    if best_node and best_len >= 400:
        return best_node

    # 2) Fallback: choose the element with the largest text density
    # (avoid picking body if it’s mostly chrome by stripping first)
    for node in soup.find_all(["div", "section", "article", "main"], recursive=True):
        txt = node.get_text("\n", strip=True)
        ln = len(txt)
        if ln > best_len:
            best_len = ln
            best_node = node

    return best_node if best_node else soup.body or soup


def extract_main_text_from_html(html: str) -> tuple[str, str]:
    """
    Returns (title, main_text).
    Uses BeautifulSoup heuristics; optionally trafilatura if installed.
    """
    # Optional: trafilatura can do very good main-content extraction
    try:
        import trafilatura
        extracted = trafilatura.extract(html, include_comments=False, include_tables=True)
        if extracted and len(extracted.strip()) > 300:
            # Title is not always available via trafilatura; we’ll still parse for title below.
            pass
        else:
            extracted = None
    except Exception:
        extracted = None

    soup = BeautifulSoup(html, "lxml")
    title = ""
    if soup.title and soup.title.get_text(strip=True):
        title = soup.title.get_text(strip=True)

    soup_remove_useless(soup)

    if extracted:
        main_text = extracted.strip()
        return title, main_text

    container = pick_best_container(soup)
    raw_text = container.get_text("\n", strip=True)
    lines = normalize_lines(raw_text)
    main_text = "\n".join(lines).strip()
    return title, main_text


# -----------------------------
# WORD OUTPUT
# -----------------------------

def save_to_docx(docx_path: Path, title: str, main_text: str, source_url: str) -> None:
    doc = Document()

    # Basic styling
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=1)
    else:
        doc.add_heading("Knowledge Article", level=1)

    doc.add_paragraph(f"Source: {source_url}")

    doc.add_paragraph("")  # spacer

    # Preserve simple structure: paragraphs and bullets
    for line in main_text.splitlines():
        l = line.strip()
        if not l:
            continue

        # Basic bullet detection
        if re.match(r"^(\-|\*|•)\s+", l):
            doc.add_paragraph(re.sub(r"^(\-|\*|•)\s+", "", l), style="List Bullet")
        elif re.match(r"^\d+[\.\)]\s+", l):
            doc.add_paragraph(re.sub(r"^\d+[\.\)]\s+", "", l), style="List Number")
        else:
            # Heading-ish heuristic
            if len(l) <= 80 and (l.isupper() or l.endswith(":")):
                doc.add_heading(l.rstrip(":"), level=2)
            else:
                doc.add_paragraph(l)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--url", default=DEFAULT_URL)
    ap.add_argument("--outdir", default="downloads")
    ap.add_argument("--profile-dir", default="pw_profile")
    ap.add_argument("--relogin", action="store_true")
    ap.add_argument("--headed", action="store_true")
    ap.add_argument("--timeout", type=int, default=240_000)
    ap.add_argument("--min-chars", type=int, default=800)
    ap.add_argument("--name", default="")

    # NEW: Word output options
    ap.add_argument("--docx", action="store_true", help="Save main content to a Word .docx file")
    ap.add_argument("--docx-name", default="", help="Docx filename (optional)")

    # Sometimes the nav wrapper works better; keep option
    ap.add_argument("--skip-direct", action="store_true")

    args = ap.parse_args()

    out_dir = Path(args.outdir)
    out_dir.mkdir(parents=True, exist_ok=True)

    profile_dir = Path(args.profile_dir)
    if args.relogin and profile_dir.exists():
        print(f"🧹 Removing profile dir for relogin: {profile_dir}")
        shutil.rmtree(profile_dir, ignore_errors=True)

    direct_url = decode_target_to_direct_url(args.url)
    kb = extract_kb_number(direct_url) or extract_kb_number(args.url) or "KB"
    base_name = safe_filename(args.name if args.name.strip() else kb)
    base = f"{base_name}_{stamp()}"

    with sync_playwright() as p:
        context = p.chromium.launch_persistent_context(
            user_data_dir=str(profile_dir),
            headless=(not args.headed),
            viewport={"width": 1400, "height": 900},
        )
        page = context.new_page()
        page.set_default_timeout(args.timeout)
        attach_debug(page)

        # Open initial URL
        print("➡️ Opening initial URL...")
        goto_with_sso_retry(page, args.url, timeout_ms=args.timeout, headed=args.headed)

        # Manual SSO if needed
        if looks_like_login(page.url):
            if not args.headed:
                context.close()
                raise RuntimeError("SSO requires interaction. Re-run with --headed.")
            input("✅ Complete SSO in the browser, then press ENTER here...")

        # Navigate to direct KB URL unless skipped
        if not args.skip_direct:
            print("➡️ Navigating to direct KB URL...")
            goto_with_sso_retry(page, direct_url, timeout_ms=args.timeout, headed=args.headed)

        # Wait for content
        print(f"⏳ Waiting for body text length >= {args.min_chars}")
        length = wait_for_text_length(page, min_chars=args.min_chars, timeout_ms=args.timeout)
        print(f"   Observed text length: {length}")

        # Save raw HTML + screenshot (still useful)
        png_path = out_dir / f"{base}.png"
        html_path = out_dir / f"{base}.html"
        try:
            page.screenshot(path=str(png_path), full_page=True)
            print(f"✅ Screenshot: {png_path}")
        except Exception as e:
            print(f"⚠️ Screenshot failed: {e}")

        html = page.content()
        html_path.write_text(html, encoding="utf-8")
        print(f"✅ HTML: {html_path}")

        # NEW: Extract main content and write Word file
        if args.docx:
            title, main_text = extract_main_text_from_html(html)

            if not main_text or len(main_text) < 200:
                print("⚠️ Main content extraction produced little text.")
                print("   Tip: the page may be mostly dynamic; increase --timeout or --min-chars, or share HTML for selector tuning.")

            docx_filename = args.docx_name.strip() or f"{base}.docx"
            docx_path = out_dir / safe_filename(docx_filename, default=f"{base}.docx")
            save_to_docx(docx_path, title=title, main_text=main_text, source_url=page.url)

            print(f"✅ Word saved: {docx_path}")

        context.close()
        print("\nDone.")


if __name__ == "__main__":
    main()
