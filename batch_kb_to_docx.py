#!/usr/bin/env python3
import argparse
import base64
import re
import shutil
import time
import tempfile
from io import BytesIO
from pathlib import Path
from datetime import datetime
from urllib.parse import unquote, urljoin, urlparse

import pandas as pd
from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from docx.shared import Inches, Pt

from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError
from playwright._impl._errors import Error as PlaywrightError


# -------------------------
# Helpers
# -------------------------

def stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def safe_filename(name: str, default: str = "file") -> str:
    name = (name or "").strip() or default
    name = re.sub(r"[^\w\-\.]+", "_", name)
    return name[:180]


def looks_like_login(url: str) -> bool:
    u = (url or "").lower()
    return any(s in u for s in (
        "login", "sso", "saml", "auth", "okta", "adfs", "signin", "microsoftonline.com"
    ))


def extract_kb_number(text: str) -> str | None:
    m = re.search(r"(KB\d+)", text or "")
    return m.group(1) if m else None


def decode_target_to_direct_url(url: str) -> str:
    """
    Converts:
      .../target/kb_view.do%3Fsysparm_article%3DKB0010611
    into:
      https://<host>/kb_view.do?sysparm_article=KB0010611
    """
    m = re.search(r"/target/([^?]+)", url)
    if not m:
        return url
    encoded = m.group(1)
    decoded = unquote(encoded)
    if decoded.startswith("http://") or decoded.startswith("https://"):
        return decoded
    host = re.match(r"^(https?://[^/]+)", url)
    if not host:
        return url
    return f"{host.group(1)}/{decoded.lstrip('/')}"


def attach_debug(page, verbose: bool = False):
    if not verbose:
        return
    page.on("console", lambda msg: print(f"[console] {msg.type}: {msg.text}"))
    page.on("pageerror", lambda err: print(f"[pageerror] {err}"))
    page.on("requestfailed", lambda req: print(f"[requestfailed] {req.url} -> {req.failure}"))


def get_body_text(page) -> str:
    try:
        return page.evaluate("() => (document.body && document.body.innerText || '').trim()")
    except Exception:
        return ""


def wait_for_text_length(page, min_chars: int, timeout_ms: int) -> int:
    deadline = time.time() + timeout_ms / 1000
    last_len = 0
    while time.time() < deadline:
        txt = get_body_text(page)
        last_len = len(txt)
        if last_len >= min_chars:
            return last_len
        time.sleep(0.25)
    return last_len


def goto_with_sso_retry(page, url: str, timeout_ms: int, headed: bool, max_attempts: int = 3) -> None:
    """
    Navigate and tolerate SSO interruptions (Azure AD hops).
    """
    for attempt in range(1, max_attempts + 1):
        try:
            page.goto(url, wait_until="domcontentloaded", timeout=timeout_ms)
            return
        except PlaywrightTimeoutError:
            # Page might still be usable; proceed
            return
        except PlaywrightError as e:
            msg = str(e)
            if "is interrupted by another navigation" in msg:
                # Let the other navigation finish (often SSO)
                try:
                    page.wait_for_load_state("domcontentloaded", timeout=30_000)
                except Exception:
                    pass

                if looks_like_login(page.url):
                    if not headed:
                        raise RuntimeError(
                            "SSO requires interaction but you are running headless. Re-run with --headed."
                        )
                    input("✅ Complete SSO in the browser window, then press ENTER to retry...")
                continue
            raise


# -------------------------
# Main content extraction (HTML -> container Tag)
# -------------------------

USELESS_TAGS = {"script", "style", "noscript", "svg", "canvas"}
USELESS_SELECTORS = [
    "header", "footer", "nav",
    "[role='navigation']", "[role='banner']", "[role='contentinfo']",
    ".navbar", ".nav", ".navigation", ".header", ".footer",
    ".sidebar", ".side-nav", ".toc", ".breadcrumbs",
]


def soup_remove_useless(soup: BeautifulSoup) -> None:
    for t in list(soup.find_all(list(USELESS_TAGS))):
        t.decompose()
    for sel in USELESS_SELECTORS:
        for t in soup.select(sel):
            t.decompose()


def pick_best_container(soup: BeautifulSoup) -> Tag:
    candidate_selectors = [
        "article", "main",
        "#kb_article", ".kb-article", ".kb-article-content", ".kb-view",
        ".sn-kb-article", ".knowledge-article",
        "[id*='kb']", "[class*='kb']",
        "[class*='article']", "[class*='content']",
    ]

    best_node = None
    best_len = 0

    for sel in candidate_selectors:
        for node in soup.select(sel):
            txt = node.get_text("\n", strip=True)
            ln = len(txt)
            if ln > best_len:
                best_len = ln
                best_node = node

    if best_node and best_len >= 400:
        return best_node

    for node in soup.find_all(["div", "section", "article", "main"], recursive=True):
        txt = node.get_text("\n", strip=True)
        ln = len(txt)
        if ln > best_len:
            best_len = ln
            best_node = node

    return best_node if best_node else (soup.body or soup)


def extract_main_container(html: str) -> tuple[str, Tag]:
    soup = BeautifulSoup(html, "lxml")
    title = soup.title.get_text(strip=True) if soup.title else ""
    soup_remove_useless(soup)
    container = pick_best_container(soup)
    return title, container


# -------------------------
# Images (auth-aware via context.request)
# -------------------------

def is_data_image(src: str) -> bool:
    return src.startswith("data:image/")


def decode_data_image(src: str) -> tuple[bytes, str]:
    header, b64 = src.split(",", 1)
    m = re.search(r"data:image/([^;]+);base64", header, re.IGNORECASE)
    ext = (m.group(1).lower() if m else "png").replace("jpeg", "jpg")
    return base64.b64decode(b64), ext


def guess_ext_from_content_type(ct: str | None) -> str:
    if not ct:
        return "png"
    ct = ct.lower()
    if "jpeg" in ct:
        return "jpg"
    if "png" in ct:
        return "png"
    if "gif" in ct:
        return "gif"
    if "webp" in ct:
        return "webp"
    if "bmp" in ct:
        return "bmp"
    if "tiff" in ct:
        return "tif"
    return "png"


def fetch_image_bytes(page, context, src: str, base_url: str) -> tuple[bytes | None, str | None]:
    if not src:
        return None, None

    src = src.strip()

    # data: image
    if is_data_image(src):
        bts, ext = decode_data_image(src)
        return bts, ext

    # blob: image -> fetch within page
    if src.startswith("blob:"):
        try:
            b64 = page.evaluate(
                """async (u) => {
                    const r = await fetch(u);
                    const b = await r.blob();
                    const a = new Uint8Array(await b.arrayBuffer());
                    let s = '';
                    for (let i=0; i<a.length; i++) s += String.fromCharCode(a[i]);
                    return btoa(s);
                }""",
                src
            )
            return base64.b64decode(b64), "png"
        except Exception:
            return None, None

    full_url = urljoin(base_url, src)
    scheme = urlparse(full_url).scheme.lower()
    if scheme not in ("http", "https"):
        return None, None

    try:
        resp = context.request.get(full_url, timeout=60_000)
        if not resp.ok:
            return None, None
        ext = guess_ext_from_content_type(resp.headers.get("content-type"))
        return resp.body(), ext
    except Exception:
        return None, None


# -------------------------
# HTML -> DOCX (structure + images)
# -------------------------

def build_docx(title: str, source_url: str, container: Tag, page, context,
               output_path: Path, max_image_width_in: float = 6.0):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title or "Knowledge Article", level=1)
    doc.add_paragraph(f"Source: {source_url}")
    doc.add_paragraph("")

    temp_dir = Path(tempfile.mkdtemp(prefix="kb_docx_imgs_"))
    img_counter = 0

    def add_text_paragraph(text: str):
        text = re.sub(r"\s+", " ", (text or "")).strip()
        if text:
            doc.add_paragraph(text)

    def add_heading(text: str, level: int):
        text = re.sub(r"\s+", " ", (text or "")).strip()
        if text:
            doc.add_heading(text, level=level)

    def add_list_item(text: str, ordered: bool):
        text = re.sub(r"\s+", " ", (text or "")).strip()
        if not text:
            return
        doc.add_paragraph(text, style=("List Number" if ordered else "List Bullet"))

    def handle_img(img_tag: Tag):
        nonlocal img_counter
        src = img_tag.get("src") or img_tag.get("data-src") or ""
        alt = img_tag.get("alt") or ""

        bts, ext = fetch_image_bytes(page, context, src, base_url=source_url)
        if not bts:
            # If it fails, add a note so you know something was there
            doc.add_paragraph(f"[Image not downloaded]{(': ' + alt) if alt else ''}")
            return

        img_counter += 1
        img_path = temp_dir / f"img_{img_counter}.{ext}"
        img_path.write_bytes(bts)

        if alt:
            doc.add_paragraph(alt)

        try:
            doc.add_picture(str(img_path), width=Inches(max_image_width_in))
        except Exception:
            # Word/python-docx may not support webp etc.
            doc.add_paragraph(f"[Image format not supported in Word: .{ext}]")

    def walk(node: Tag):
        if isinstance(node, NavigableString):
            return
        if not isinstance(node, Tag):
            return

        name = (node.name or "").lower()

        if name in ("script", "style", "noscript"):
            return

        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = min(4, int(name[1]))
            add_heading(node.get_text(" ", strip=True), level=level)
            return

        if name == "p":
            txt = node.get_text(" ", strip=True)
            if txt:
                add_text_paragraph(txt)
            for img in node.find_all("img"):
                handle_img(img)
            return

        if name in ("ul", "ol"):
            ordered = (name == "ol")
            for li in node.find_all("li", recursive=False):
                add_list_item(li.get_text(" ", strip=True), ordered=ordered)
                for img in li.find_all("img"):
                    handle_img(img)
            return

        if name == "img":
            handle_img(node)
            return

        if name == "table":
            # Simple fallback: table to text
            txt = node.get_text("\n", strip=True)
            if txt:
                add_text_paragraph(txt)
            for img in node.find_all("img"):
                handle_img(img)
            return

        # default: recurse direct children to keep order
        for child in node.children:
            if isinstance(child, Tag):
                walk(child)

    for child in container.children:
        if isinstance(child, Tag):
            walk(child)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # Cleanup temp image dir
    try:
        shutil.rmtree(temp_dir, ignore_errors=True)
    except Exception:
        pass


# -------------------------
# Batch: Excel reading + crawling loop
# -------------------------

def read_targets_from_excel(excel_path: Path, sheet: str | int | None,
                            kb_col: str | None, url_col: str | None):
    df = pd.read_excel(excel_path, sheet_name=sheet, engine="openpyxl")
    df.columns = [str(c).strip() for c in df.columns]

    kb_col = kb_col or ("KB" if "KB" in df.columns else None)
    url_col = url_col or ("URL" if "URL" in df.columns else None)

    if not kb_col and not url_col:
        raise ValueError(
            f"Excel must contain a '{'KB'}' column or a '{'URL'}' column. Found: {list(df.columns)}"
        )

    targets = []
    for idx, row in df.iterrows():
        kb = str(row.get(kb_col, "")).strip() if kb_col else ""
        url = str(row.get(url_col, "")).strip() if url_col else ""

        kb = kb if kb and kb.lower() != "nan" else ""
        url = url if url and url.lower() != "nan" else ""

        if not kb and not url:
            continue

        if kb:
            kb_match = extract_kb_number(kb)
            if kb_match:
                kb = kb_match
            else:
                # If kb column contains a URL, extract KB from it
                kb2 = extract_kb_number(kb)
                kb = kb2 or kb

        if not kb and url:
            kb = extract_kb_number(url) or ""

        targets.append({"row": int(idx) + 2, "kb": kb, "url": url})
    return targets, df


def main():
    ap = argparse.ArgumentParser(description="Batch crawl ServiceNow KBs from Excel and export to DOCX (with images).")
    ap.add_argument("--excel", required=True, help="Path to Excel file (.xlsx)")
    ap.add_argument("--sheet", default=0, help="Sheet name or index (default: 0)")
    ap.add_argument("--kb-col", default="", help="Column name containing KB numbers (default: KB)")
    ap.add_argument("--url-col", default="", help="Column name containing full URLs (default: URL)")
    ap.add_argument("--outdir", default="kb_exports", help="Output folder for docx files")
    ap.add_argument("--profile-dir", default="pw_profile", help="Persistent Playwright profile directory")
    ap.add_argument("--relogin", action="store_true", help="Delete profile dir and login again")
    ap.add_argument("--headed", action="store_true", help="Run with a visible browser (recommended for SSO)")
    ap.add_argument("--timeout", type=int, default=240_000, help="Timeout ms per KB (default 240000)")
    ap.add_argument("--min-chars", type=int, default=800, help="Minimum text length to treat page as loaded")
    ap.add_argument("--base-host", default="", help="Base host, e.g. https://myservice.lloyds.com (optional)")
    ap.add_argument("--url-template", default="/kb_view.do?sysparm_article={KB}",
                    help="Template for KB URL if only KB number is given")
    ap.add_argument("--skip-direct", action="store_true",
                    help="Do not decode /target/ url into direct kb_view.do url")
    ap.add_argument("--sleep", type=float, default=1.0, help="Seconds to sleep between KBs (default 1.0)")
    ap.add_argument("--max", type=int, default=0, help="Max KBs to process (0 = all)")
    ap.add_argument("--verbose", action="store_true", help="Verbose Playwright console/request logging")
    args = ap.parse_args()

    excel_path = Path(args.excel)
    out_dir = Path(args.outdir)
    out_dir.mkdir(parents=True, exist_ok=True)

    profile_dir = Path(args.profile_dir)
    if args.relogin and profile_dir.exists():
        print(f"🧹 Removing profile dir for relogin: {profile_dir}")
        shutil.rmtree(profile_dir, ignore_errors=True)

    sheet = args.sheet
    try:
        sheet = int(sheet)
    except Exception:
        pass

    kb_col = args.kb_col.strip() or None
    url_col = args.url_col.strip() or None

    targets, _ = read_targets_from_excel(excel_path, sheet=sheet, kb_col=kb_col, url_col=url_col)
    if args.max and args.max > 0:
        targets = targets[:args.max]

    if not targets:
        print("No targets found in Excel.")
        return

    base_host = args.base_host.strip().rstrip("/")
    url_template = args.url_template

    results = []

    with sync_playwright() as p:
        context = p.chromium.launch_persistent_context(
            user_data_dir=str(profile_dir),
            headless=(not args.headed),
            viewport={"width": 1400, "height": 900},
        )
        page = context.new_page()
        page.set_default_timeout(args.timeout)
        attach_debug(page, verbose=args.verbose)

        # Open a harmless page to establish context (optional)
        if base_host:
            try:
                goto_with_sso_retry(page, base_host, timeout_ms=args.timeout, headed=args.headed)
            except Exception:
                pass

        # One-time login if we land on SSO
        if looks_like_login(page.url):
            if not args.headed:
                context.close()
                raise RuntimeError("SSO requires interaction. Re-run with --headed.")
            input("✅ Complete SSO in the browser, then press ENTER to begin batch...")

        for i, t in enumerate(targets, start=1):
            rownum = t["row"]
            kb = (t["kb"] or "").strip()
            url = (t["url"] or "").strip()

            # Build URL if only KB number provided
            if not url and kb:
                if not base_host:
                    raise ValueError("You provided KB numbers but no --base-host. Example: --base-host https://myservice.lloyds.com")
                url = base_host + url_template.format(KB=kb)

            # If URL is the /target/ wrapper, optionally decode to direct kb_view
            if url and (not args.skip_direct):
                url = decode_target_to_direct_url(url)

            # Determine output filename (KB number preferred)
            kb_for_name = extract_kb_number(kb) or extract_kb_number(url) or f"ROW{rownum}"
            outfile = out_dir / f"{safe_filename(kb_for_name)}.docx"

            print(f"\n[{i}/{len(targets)}] Row {rownum} -> {kb_for_name} -> {url}")
            started = time.time()
            status = "OK"
            error = ""
            final_url = ""
            title = ""

            try:
                goto_with_sso_retry(page, url, timeout_ms=args.timeout, headed=args.headed)

                # If redirected to login mid-batch, let user fix once and continue
                if looks_like_login(page.url):
                    if not args.headed:
                        raise RuntimeError("Redirected to SSO in headless mode. Re-run with --headed.")
                    input("🔐 SSO appeared again. Complete login, then press ENTER to continue...")
                    goto_with_sso_retry(page, url, timeout_ms=args.timeout, headed=args.headed)

                # Wait for content
                observed = wait_for_text_length(page, min_chars=args.min_chars, timeout_ms=args.timeout)
                final_url = page.url

                # Save DOCX (main content + images)
                html = page.content()
                title, container = extract_main_container(html)
                build_docx(title=title, source_url=final_url, container=container, page=page, context=context,
                           output_path=outfile, max_image_width_in=6.0)

                if observed < max(200, args.min_chars // 2):
                    status = "WARN_THIN_CONTENT"

            except Exception as e:
                status = "FAIL"
                error = str(e)[:2000]
                print(f"❌ Failed: {error}")

            elapsed = round(time.time() - started, 2)
            results.append({
                "row": rownum,
                "kb": kb_for_name,
                "input_url": t["url"] or "",
                "used_url": url,
                "final_url": final_url,
                "title": title,
                "outfile": str(outfile),
                "status": status,
                "elapsed_sec": elapsed,
                "error": error
            })

            time.sleep(max(0.0, args.sleep))

        context.close()

    # Write results log
    results_df = pd.DataFrame(results)
    log_path = out_dir / f"results_{stamp()}.xlsx"
    results_df.to_excel(log_path, index=False, engine="openpyxl")
    print(f"\n✅ Batch complete. Log written to: {log_path}")


if __name__ == "__main__":
    main()
